# coding=utf-8
# 83条记录
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE


# https://python-docx.readthedocs.io/en/latest/user/styles-understanding.html#understanding-styles
# http://www.cnblogs.com/rencm/p/6285304.html
# https://zhuanlan.zhihu.com/p/21716087
# http://blog.csdn.net/qianchenglenger/article/details/51582005 设置字体部分有效

class HtmlOutputer(object):
    def __init__(self):
        self.datas = []

    def collect_data(self, new_data):
        if new_data is None:
            return
        self.datas.append(new_data)

    def out_html(self):
        """
        fout = open('output.txt', 'w', encoding='UTF-8')
        fout.write("<html>")
        fout.write("<head><meta http-equiv=\"content-type\" content=\"text/html;charset=utf-8\"></head>")
        fout.write("<body>")
        fout.write("<table>")

        for data in self.datas:
            fout.write("<tr>")
            fout.write("<td>%s</td>" % data['url'])
            fout.write("<td>%s</td>" % data['title'])
            fout.write("<td>%s</td>" % data['summary'])
            fout.write("<tr>")

        fout.write("</table>")
        fout.write("</body>")
        fout.write("</html>")
        fout.close()
        """

        """
        doc = Document()
        # styles = doc.styles
        # print("\"\\n\".join([s.name for s in styles if s.type == WD_STYLE_TYPE.PARAGRAPH])")
        for data in self.datas:
            doc.add_heading(data['title'])

            paragraph = doc.add_paragraph(data['summary'])

            run = paragraph.add_run(u'设置字号、')
            run.font.size = Pt(24)

            run = paragraph.add_run(u'设置中文字体、')
            run.font.name = u'宋体'
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            doc.save("demo.docx")
        """
        fout = open('output.txt', 'w', encoding='UTF-8')

        for data in self.datas:
            fout.write(data['title'] + "\n")
            fout.write(data['summary'] + "\n")
        fout.close()
